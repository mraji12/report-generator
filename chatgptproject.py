import streamlit as st
from streamlit_option_menu import option_menu
from google import genai
from docx import Document
import io
import os
0
selected = option_menu(
    menu_title=None,
    options=["Home", "Sign in", "Log in"],
    icons=["house-fill", "door-closed-fill", "door-open-fill"],
    default_index=0,
    orientation="horizontal",
    styles={
        "container": {"padding": "0!important", "background-color": " "},
        "icon": {"color": "orange", "font-size": "15px"},
        "nav-link": {"font-size": "15px", "text-align": "center", "margin": "0px", "--hover-color": "#3944BC"},
        "nav-link-selected": {"background-color": "#3944BC "},
    }
)

client = genai.Client(api_key="AIzaSyDf1hVPWes0ydntTM9lRsfmpxOQkMoi7SY")
model = client.chats.create(model="gemini-2.0-flash")

st.title("Project Report Generator")

# Show Figure 5.3 if image file is present
figure_path = os.path.join(os.path.dirname(__file__), "assets", "figure_5_3.png")
if os.path.exists(figure_path):
    st.image(figure_path, caption="Figure 5.3 – Component Structure in VS Code", use_column_width=True)

uploaded_file = st.file_uploader("Upload your project document (.docx only)", type=["docx"])

if uploaded_file is not None:
    try:
        doc = Document(io.BytesIO(uploaded_file.read()))
        doc_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip() != ""])
        st.subheader("Generate Project Report")
        if st.button("Generate Report"):
            with st.spinner("Generating your project report..."):
                prompt = (
                    "You are a project report assistant. Read the following project document and generate a well-structured project report summary with introduction, objectives, methodology, results, and conclusion sections:\n\n"
                    + doc_text
                )
                response = model.send_message(prompt)
                st.subheader("Generated Project Report:")
                st.write(response.text)

                # --- Download as .docx ---
                output_doc = Document()
                output_doc.add_heading("Generated Project Report", 0)
                for para in response.text.split('\n'):
                    if para.strip():
                        output_doc.add_paragraph(para.strip())
                docx_io = io.BytesIO()
                output_doc.save(docx_io)
                docx_io.seek(0)
                st.download_button(
                    label="Download Report as .docx",
                    data=docx_io,
                    file_name="project_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    except Exception as e:
        if "overloaded" in str(e).lower() or "503" in str(e):
            st.error("The AI model is overloaded. Please try again in a few minutes.")
        else:
            st.error(f"Error: {e}")
else:
    st.info("Please upload a .docx file to generate your project report.")